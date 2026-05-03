"""
services.py — Core business logic for YouTube Video Summariser.

Contains all the heavy-lifting functions:
  - Transcript extraction from YouTube
  - Summary generation via Google Gemini
  - Mermaid.js diagram code generation via Gemini
  - Word document (.docx) creation
"""

import os
import re
import io
from youtube_transcript_api import YouTubeTranscriptApi
from google import genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ──────────────────────────────────────────────
# 1. TRANSCRIPT EXTRACTION
# ──────────────────────────────────────────────

def extract_video_id(url: str) -> str:
    """
    Extracts the YouTube video ID from various URL formats:
      - https://www.youtube.com/watch?v=VIDEO_ID
      - https://youtu.be/VIDEO_ID
      - https://www.youtube.com/embed/VIDEO_ID
    """
    patterns = [
        r'(?:v=)([a-zA-Z0-9_-]{11})',
        r'(?:youtu\.be/)([a-zA-Z0-9_-]{11})',
        r'(?:embed/)([a-zA-Z0-9_-]{11})',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    raise ValueError(f"Could not extract video ID from URL: {url}")


def get_video_transcript(video_url: str) -> str:
    """
    Fetches the transcript/subtitles of a YouTube video.
    Returns the full transcript as a single string.
    """
    try:
        video_id = extract_video_id(video_url)
        api = YouTubeTranscriptApi()
        transcript = api.fetch(video_id)
        # Combine all text snippets into one block of text
        full_text = " ".join([snippet.text for snippet in transcript.snippets])
        return full_text
    except Exception as e:
        raise ValueError(f"Could not retrieve transcript for this video. Make sure it has closed captions. Details: {str(e)}")


# ──────────────────────────────────────────────
# 2. GEMINI AI — SUMMARY GENERATION
# ──────────────────────────────────────────────

def _get_gemini_client():
    """Creates and returns a configured Gemini client."""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key or api_key == "PASTE_YOUR_KEY_HERE":
        raise ValueError("GEMINI_API_KEY is not set. Please add it to your .env file.")
    client = genai.Client(api_key=api_key)
    return client


def generate_summary(transcript: str, token_limit: int = 500) -> dict:
    """
    Sends the transcript to Google Gemini and asks for a structured summary.
    Returns a dict with 'title' and 'summary' keys.
    """
    client = _get_gemini_client()

    prompt = f"""You are an expert content analyst. Summarize the following YouTube video transcript.

RULES:
1. Keep the summary under {token_limit} words.
2. Start with a concise, descriptive TITLE for the video (infer from content).
3. Structure the summary with clear sections using markdown headings (##).
4. Use bullet points for key takeaways.
5. Highlight important terms in **bold**.
6. End with a "Key Takeaways" section listing the top 3-5 insights.

FORMAT YOUR RESPONSE EXACTLY LIKE THIS:
TITLE: <your inferred title>

<your structured summary in markdown>

TRANSCRIPT:
{transcript}
"""

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
    )

    response_text = response.text

    # Extract title from response
    title = "YouTube Video Summary"
    if "TITLE:" in response_text:
        title_line = response_text.split("TITLE:")[1].split("\n")[0].strip()
        title = title_line
        # Remove the TITLE line from the summary body
        response_text = response_text.split("\n", 2)[-1].strip() if response_text.count("\n") >= 2 else response_text

    return {
        "title": title,
        "summary": response_text,
    }


# ──────────────────────────────────────────────
# 3. GEMINI AI — MERMAID DIAGRAM GENERATION
# ──────────────────────────────────────────────

def generate_mermaid_code(summary: str) -> str:
    """
    Sends the summary to Google Gemini and asks it to produce
    a Mermaid.js flowchart / block diagram representing the key concepts.
    Returns the raw Mermaid syntax string.
    """
    client = _get_gemini_client()

    prompt = f"""You are a visual information designer. Convert the following summary into a Mermaid.js diagram.

RULES:
1. Use a `graph TD` (top-down) flowchart.
2. Create a clear hierarchy: Main Topic → Sub Topics → Key Details.
3. Keep labels SHORT (max 6 words per node).
4. Use different node shapes for variety:
   - Round brackets for main topics: (Main Topic)
   - Square brackets for sub-topics: [Sub Topic]
   - Curly brackets for details: {{Detail}}
5. Use descriptive arrow labels where helpful: -->|"label"|
6. Maximum 15 nodes to keep it readable.
7. Do NOT wrap the code in markdown code fences (no ```mermaid or ``` tags).
8. Return ONLY the raw Mermaid code, nothing else.

SUMMARY:
{summary}
"""

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
    )

    mermaid_code = response.text.strip()

    # Clean up if Gemini wraps it in code fences anyway
    mermaid_code = mermaid_code.replace("```mermaid", "").replace("```", "").strip()

    return mermaid_code


# ──────────────────────────────────────────────
# 4. WORD DOCUMENT (.docx) GENERATION
# ──────────────────────────────────────────────

def create_word_document(title: str, summary: str) -> io.BytesIO:
    """
    Creates a professionally styled Word document from the summary.
    Returns a BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # ── Page Styling ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51)

    # ── Title ──
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(26, 26, 64)  # Dark navy

    # ── Separator Line ──
    doc.add_paragraph("─" * 60)

    # ── Parse Markdown Summary into Structured Docx ──
    lines = summary.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Heading level 2: ## Section Title
        if stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            heading = doc.add_heading(heading_text, level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 102, 153)  # Teal

        # Heading level 3: ### Sub Section
        elif stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            heading = doc.add_heading(heading_text, level=3)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 128, 128)

        # Bullet point: - item or * item
        elif stripped.startswith(('- ', '* ')):
            bullet_text = stripped[2:].strip()
            # Handle bold markers **text**
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, bullet_text)

        # Regular paragraph
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)

    # ── Footer ──
    doc.add_paragraph()
    footer = doc.add_paragraph("─" * 60)
    footer_text = doc.add_paragraph("Generated by YouTube Video Summariser")
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_text.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

    # Save to in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_formatted_text(paragraph, text: str):
    """
    Parses simple markdown bold (**text**) within a line
    and adds runs with appropriate formatting to the paragraph.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
